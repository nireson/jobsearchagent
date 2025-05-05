import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
import json
from typing import Optional, List, Union, Dict, Any

# Import the LLM client based on environment settings
def get_llm_client():
    """Get the appropriate LLM client based on environment settings."""
    provider = os.environ.get('MODEL_PROVIDER', 'openai').lower()
    
    if provider == 'openai':
        from utils.llm_clients import OpenAIClient
        return OpenAIClient()
    elif provider == 'anthropic':
        from utils.llm_clients import AnthropicClient
        return AnthropicClient()
    elif provider == 'ollama':
        from utils.llm_clients import OllamaClient
        return OllamaClient()
    else:
        # Default to OpenAI
        from utils.llm_clients import OpenAIClient
        return OpenAIClient()

def format_agent_output(raw_content, search_query=None, format_instructions=None):
    """
    Format the raw agent output into a well-structured document.
    
    Args:
        raw_content: The raw output from the browser agent
        search_query: The original search query for context
        format_instructions: Optional formatting instructions from the user
        
    Returns:
        str: Formatted content ready for document creation
    """
    try:
        # Get the LLM client
        llm_client = get_llm_client()
        
        # Prepare prompt for the formatting agent
        system_prompt = """You are a document formatter. Your job is to take the raw output from a web search agent and format it into a clear, structured document.
        
Extract the relevant information from the search results and format them into a professional document with these requirements:
1. Create clear, concise sections with appropriate headings
2. Extract and preserve all links, maintaining them as clickable links
3. Format search results in a logical and organized manner
4. Organize information logically with bullet points where appropriate
5. Remove any debugging information, logs, or technical details
6. Focus only on the actual search results, not the search process
7. Use professional language and formatting

Return only the formatted content without explanations. The output will be used directly in a Word document."""

        # Use the user's format instructions if provided
        user_prompt = f"""Here is the search query: {search_query or "General search"}

And here is the raw output from the search agent:

{raw_content}"""

        if format_instructions and len(format_instructions.strip()) > 0:
            user_prompt += f"""

Please format this content according to these specific instructions:

{format_instructions}"""
        else:
            user_prompt += """

Please format this into a professional document."""

        # Call the LLM to format the content
        formatted_content = llm_client.chat_completion(
            system_prompt=system_prompt,
            user_prompt=user_prompt,
            max_tokens=4000
        )
        
        # If we don't get a valid response, do basic cleaning instead
        if not formatted_content or len(formatted_content) < 100:
            print("Warning: LLM formatting returned insufficient content, falling back to basic cleaning")
            return clean_content(raw_content)
            
        return formatted_content
        
    except Exception as e:
        print(f"Error formatting agent output: {str(e)}")
        # Fall back to basic cleaning if formatting fails
        return clean_content(raw_content)

def clean_content(content: Union[str, Dict, List]) -> str:
    """
    Clean and format the content for the document.
    
    Args:
        content: The raw content to clean and format. Can be a string, dictionary, or list.
        
    Returns:
        str: The cleaned and formatted content as a string.
        
    Raises:
        ValueError: If the content cannot be properly formatted.
    """
    if not content:
        return ""
        
    try:
        # If content is a string representation of an object, try to extract the actual content
        if isinstance(content, str):
            # Try to extract content from AgentHistoryList format
            if "extracted_content=" in content:
                matches = re.findall(r"extracted_content='(.*?)'", content)
                if matches:
                    # Get the last successful extraction
                    content = matches[-1]
            
            # Try to extract content from model outputs
            if "all_model_outputs" in content:
                match = re.search(r"'text': '(.*?)'", content)
                if match:
                    content = match.group(1)
        
        # Handle JSON content
        if isinstance(content, str) and (content.startswith('{') or content.startswith('[')):
            try:
                import json
                data = json.loads(content)
                if isinstance(data, dict):
                    # Format the content in a readable way
                    formatted = []
                    for key, value in data.items():
                        if isinstance(value, list):
                            formatted.append(f"\n{key.capitalize()}:")
                            for item in value:
                                if isinstance(item, dict):
                                    # Handle nested dictionaries
                                    nested = []
                                    for k, v in item.items():
                                        if isinstance(v, str):
                                            nested.append(f"{k.capitalize()}: {v}")
                                    formatted.append("- " + "; ".join(nested))
                                else:
                                    formatted.append(f"- {item}")
                        elif isinstance(value, dict):
                            # Handle nested dictionaries
                            nested = []
                            for k, v in value.items():
                                if isinstance(v, str):
                                    nested.append(f"{k.capitalize()}: {v}")
                            formatted.append(f"\n{key.capitalize()}: " + "; ".join(nested))
                        else:
                            formatted.append(f"\n{key.capitalize()}: {value}")
                    content = "\n".join(formatted)
                elif isinstance(data, list):
                    # Format list items
                    formatted = []
                    for item in data:
                        if isinstance(item, dict):
                            # Handle dictionaries in list
                            nested = []
                            for k, v in item.items():
                                if isinstance(v, str):
                                    nested.append(f"{k.capitalize()}: {v}")
                            formatted.append("- " + "; ".join(nested))
                        else:
                            formatted.append(f"- {item}")
                    content = "\n".join(formatted)
            except json.JSONDecodeError:
                # If JSON parsing fails, continue with the original content
                pass
        
        # Clean up any markdown or special characters
        content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)  # Remove bold
        content = re.sub(r'\*(.*?)\*', r'\1', content)      # Remove italic
        content = re.sub(r'`(.*?)`', r'\1', content)        # Remove code blocks
        
        # Convert markdown links to HTML links
        content = re.sub(r'\[(.*?)\]\(([^)]+)\)', r'<a href="\2">\1</a>', content)
        
        # Clean up any remaining escape characters
        content = content.replace('\\n', '\n')
        content = content.replace('\\"', '"')
        content = content.replace("\\'", "'")
        
        # Remove any emoji or special characters at the start
        content = re.sub(r'^[\U0001F000-\U0001F9FF\s]*', '', content)
        
        # Ensure proper spacing between sections
        content = re.sub(r'\n{3,}', '\n\n', content)
        
        return content.strip()
        
    except Exception as e:
        raise ValueError(f"Error cleaning content: {str(e)}")

def add_hyperlink(paragraph, url, text):
    """
    Add a hyperlink to a paragraph.
    
    Args:
        paragraph: The paragraph to add the hyperlink to
        url: The URL for the hyperlink
        text: The text to display for the hyperlink
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    hyperlink.set(qn("w:history"), "1")
    
    # Create a new run
    new_run = OxmlElement("w:r")
    
    # Set run properties
    rPr = OxmlElement("w:rPr")
    
    # Add color and underline
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0000FF")  # Blue color by default
    rPr.append(c)
    
    # Add underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    
    # Join all the xml elements together
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)
    
    return hyperlink

def save_result_file(result_content, filename=None, output_format='docx', search_query=None, format_instructions=None, execution_logs=None):
    """Save the result content as a DOCX or XLSX file.

    Args:
        result_content: The final content string to save.
        filename: The name for the output file (extension will be adjusted).
        output_format: The desired format ('docx' or 'xlsx').
        search_query: (Ignored, kept for compatibility for now)
        format_instructions: (Ignored, kept for compatibility for now)
        execution_logs: (Ignored, kept for compatibility for now)

    Returns:
        str: The path to the saved file, or None if saving failed.
    """
    results_dir = os.path.join(os.getcwd(), 'results')
    os.makedirs(results_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # Determine filename and extension
    if not filename:
        base_filename = f"output_{timestamp}"
    else:
        base_filename = os.path.splitext(filename)[0]
        
    file_extension = f".{output_format.lower()}" if output_format.lower() in ['docx', 'xlsx'] else '.docx'
    final_filename = f"{base_filename}{file_extension}"
    file_path = os.path.join(results_dir, final_filename)

    try:
        if file_extension == '.docx':
            # --- DOCX Saving Logic --- 
            document = Document()
            # Apply styles (optional but recommended)
            try:
                style = document.styles['Normal']
                font = style.font
                font.name = 'Calibri'
                font.size = Pt(11)
            except Exception as style_err:
                print(f"Warning: Could not apply default styles - {style_err}")

            # Add content with link handling
            content_to_add = str(result_content)
            md_link_regex = re.compile(r'\[(.*?)\]\(([^)]+)\)')
            url_regex = re.compile(r'(https?://[\w\.\-\/\?\=\&\%\#\~]+)') # For plain URLs

            paragraphs = content_to_add.split('\n')
            for para_text in paragraphs:
                if not para_text.strip():
                    document.add_paragraph()
                    continue
                p = document.add_paragraph()
                current_pos = 0
                md_links = list(md_link_regex.finditer(para_text))
                if not md_links:
                    plain_urls = list(url_regex.finditer(para_text))
                    if not plain_urls:
                         p.add_run(para_text)
                    else:
                         # Handle plain URLs if found
                         for match in plain_urls:
                             start, end = match.span()
                             if start > current_pos:
                                 p.add_run(para_text[current_pos:start])
                             try:
                                 add_hyperlink(p, match.group(0), match.group(0))
                             except Exception as link_err:
                                 print(f"Warning: Failed adding hyperlink {match.group(0)}: {link_err}")
                                 p.add_run(match.group(0))
                             current_pos = end
                         if current_pos < len(para_text):
                             p.add_run(para_text[current_pos:])
                else:
                    # Handle markdown links
                    for match in md_links:
                        start, end = match.span()
                        if start > current_pos:
                            p.add_run(para_text[current_pos:start])
                        try:
                            add_hyperlink(p, match.group(2), match.group(1))
                        except Exception as link_err:
                            print(f"Warning: Failed adding hyperlink {match.group(1)}: {link_err}")
                            p.add_run(match.group(1) or match.group(2))
                        current_pos = end
                    if current_pos < len(para_text):
                        p.add_run(para_text[current_pos:])
            
            document.save(file_path)
            print(f"Result saved as DOCX: {file_path}")

        elif file_extension == '.xlsx':
            # --- XLSX Saving Logic --- 
            try:
                import pandas as pd
                print(f"Attempting to save as XLSX. Content length: {len(str(result_content))}") # Log content length

                # Split content by lines and create a single-column DataFrame
                lines = str(result_content).split('\n')
                # Handle potential empty lines or whitespace lines if desired
                cleaned_lines = [line for line in lines if line.strip()] 
                if not cleaned_lines:
                    cleaned_lines = ["(No content)"] # Ensure DataFrame isn't empty
                
                df = pd.DataFrame(cleaned_lines, columns=['Results']) 
                print(f"Created DataFrame for XLSX with {len(df)} rows.")
                
                df.to_excel(file_path, index=False, engine='openpyxl')
                print(f"Result saved as XLSX: {file_path}")
            except ImportError:
                print("Error: Pandas library not found. Cannot save as XLSX. Please install pandas and openpyxl.")
                return None # Indicate failure
            except Exception as excel_err:
                print(f"Error creating DataFrame or saving as XLSX: {excel_err}")
                import traceback
                traceback.print_exc()
                return None # Indicate failure
        else:
            print(f"Error: Unsupported output format '{output_format}'")
            return None
            
        return file_path # Return path on success

    except Exception as e: # Catch errors during file processing/saving
        print(f"Error processing content or saving file: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def format_document(doc: Document) -> Document:
    """
    Apply consistent formatting to a document.
    
    Args:
        doc: The Document object to format.
        
    Returns:
        Document: The formatted Document object.
        
    Raises:
        ValueError: If the document cannot be properly formatted.
    """
    try:
        # Ensure all code below is indented under the 'try'
        # Set the default font
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)  # Add spacing after paragraphs
        
        # Configure heading styles
        for i in range(1, 5):
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.name = 'Calibri'
            heading_style.font.size = Pt(16 - i)  # Decreasing size for lower levels
            heading_style.font.bold = True
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
            heading_style.paragraph_format.keep_with_next = True  # Keep headings with subsequent text
        
        # Configure list styles
        list_bullet = doc.styles['List Bullet']
        list_bullet.font.name = 'Calibri'
        list_bullet.font.size = Pt(11)
        list_bullet.paragraph_format.left_indent = Inches(0.25)
        list_bullet.paragraph_format.space_after = Pt(6)
        list_bullet.paragraph_format.first_line_indent = Inches(-0.25)  # Proper bullet indentation
        
        list_number = doc.styles['List Number']
        list_number.font.name = 'Calibri'
        list_number.font.size = Pt(11)
        list_number.paragraph_format.left_indent = Inches(0.25)
        list_number.paragraph_format.space_after = Pt(6)
        list_number.paragraph_format.first_line_indent = Inches(-0.25)  # Proper number indentation
        
        # Configure hyperlink style
        try:
            hyperlink_style = doc.styles['Hyperlink']
        except KeyError:
            # Create the hyperlink style if it doesn't exist
            hyperlink_style = doc.styles.add_style('Hyperlink', WD_STYLE_TYPE.CHARACTER)
        
        hyperlink_style.font.color.rgb = RGBColor(0x0D, 0x6E, 0xFD)  # Bootstrap primary blue
        hyperlink_style.font.underline = True
    
        # Set the margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            # Set page size to A4
            section.page_height = Inches(11.69)
            section.page_width = Inches(8.27)
        
        return doc
            
    except Exception as e:
        raise ValueError(f"Error formatting document: {str(e)}")
